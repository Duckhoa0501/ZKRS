import os
import re
import sys

# Đảm bảo cài đặt python-docx nếu chưa có
try:
    import docx
except ImportError:
    print("Đang tự động cài đặt thư viện python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- CÁC THIẾT LẬP ĐỊNH DẠNG BẢNG CHUYÊN NGHIỆP ---
def set_cell_background(cell, fill_hex):
    """Đặt màu nền cho ô trong bảng (Hex Color)"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Đặt khoảng cách đệm (Padding) cho ô trong bảng"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def format_bold_text(paragraph, text):
    """Hỗ trợ phân tích cú pháp chữ in đậm **text** của Markdown sang Word"""
    text = text.replace("$$", "").replace("$", "") # Xóa ký hiệu công thức LaTeX thô
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Times New Roman'

def process_table(doc, table_rows):
    """Đọc dữ liệu bảng từ Markdown và khởi tạo bảng trong Word"""
    rows_data = []
    for r in table_rows:
        # Bỏ qua các dòng phân cách dạng |---|---|
        if re.match(r'^\|[\s:-|]*\|$', r):
            continue
        cells = [c.strip() for c in r.split('|')[1:-1]]
        rows_data.append(cells)
        
    if not rows_data:
        return
        
    num_cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = 'Table Grid'
    
    for r_idx, row_cells in enumerate(rows_data):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_cells):
            # Ngăn lỗi tràn cột
            if c_idx >= num_cols:
                continue
            cell = row.cells[c_idx]
            cell.text = "" # Xóa text mặc định
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            
            format_bold_text(p, cell_text)
            
            # Định dạng hàng tiêu đề (Header Row)
            if r_idx == 0:
                set_cell_background(cell, 'EAEAEA') # Nền xám nhạt sang trọng
                for run in p.runs:
                    run.bold = True
            
            # Đặt padding
            set_cell_margins(cell, top=120, bottom=120, left=180, right=180)

def convert_md_to_docx(doc, md_text):
    """Duyệt và phân tích cú pháp tệp Markdown để ghi vào Document"""
    lines = md_text.split('\n')
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Nhóm các dòng bảng Markdown
        if line.startswith('|'):
            in_table = True
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            process_table(doc, table_rows)
            in_table = False
            table_rows = []
            
        if not line:
            # Tạo khoảng trắng giãn cách đoạn
            i += 1
            continue
            
        # Heading 1
        if line.startswith('# '):
            p = doc.add_heading(level=1)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[2:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 51, 102) # Xanh dương đậm học thuật
            run.bold = True
            
        # Heading 2
        elif line.startswith('## '):
            p = doc.add_heading(level=2)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[3:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.bold = True
            
        # Heading 3
        elif line.startswith('### '):
            p = doc.add_heading(level=3)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line[4:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(51, 51, 51)
            run.bold = True
            
        # Bullet list
        elif line.startswith('* ') or line.startswith('- '):
            text = line[2:]
            # Lọc bỏ các dấu checkbox
            if text.startswith('[ ]') or text.startswith('[x]') or text.startswith('[/]'):
                text = text[3:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.left_indent = Inches(0.25)
            format_bold_text(p, text)
            
        # Trích dẫn / Hộp cảnh báo
        elif line.startswith('> '):
            text = line[2:]
            if text.startswith('[!'): # Bỏ qua tag alert của GitHub markdown
                i += 1
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.italic = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(102, 102, 102)
            
        # Khối mã nguồn (Code Blocks)
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 102, 51) # Màu xanh lục code
            
        # Đoạn văn thường
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5 # Giãn dòng 1.5 lines chuẩn đồ án
            format_bold_text(p, line)
            
        i += 1
        
    if in_table:
        process_table(doc, table_rows)

def main():
    print("--- KHỞI TẠO TIẾN TRÌNH BIÊN DỊCH BÁO CÁO DOCX ---")
    
    # Tạo đối tượng Document mới
    doc = Document()
    
    # Thiết lập căn lề chuẩn đồ án tốt nghiệp
    # Lề trái: 3cm (1.18 in), Lề phải: 2cm (0.79 in), Lề trên: 2.5cm (0.98 in), Lề dưới: 2.5cm (0.98 in)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.98)
        section.bottom_margin = Inches(0.98)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.79)

    # Đọc tệp Chương 3
    ch3_path = "C:\\Users\\admin\\.gemini\\antigravity\\brain\\cb9465fa-1c70-496e-8e78-602df853df06\\chapter_3_implementation_report.md"
    print(f"Đọc tệp Chương 3 tại: {ch3_path}")
    if os.path.exists(ch3_path):
        with open(ch3_path, "r", encoding="utf-8") as f:
            ch3_content = f.read()
        convert_md_to_docx(doc, ch3_content)
        # Thêm dấu ngắt trang giữa 2 chương
        doc.add_page_break()
    else:
        print("Lỗi: Không tìm thấy file Chương 3 md.")

    # Đọc tệp Chương 4
    ch4_path = "C:\\Users\\admin\\.gemini\\antigravity\\brain\\cb9465fa-1c70-496e-8e78-602df853df06\\chapter_4_testing_performance.md"
    print(f"Đọc tệp Chương 4 tại: {ch4_path}")
    if os.path.exists(ch4_path):
        with open(ch4_path, "r", encoding="utf-8") as f:
            ch4_content = f.read()
        convert_md_to_docx(doc, ch4_content)
    else:
        print("Lỗi: Không tìm thấy file Chương 4 md.")

    # Lưu tệp Word kết quả
    output_filename = "Bao_Cao_Chuong_3_4_Mat_Ma_E2EE.docx"
    print(f"Đang lưu file Word kết quả: {output_filename}")
    doc.save(output_filename)
    print("--- BIÊN DỊCH BÁO CÁO THÀNH CÔNG! ---")

if __name__ == "__main__":
    main()
